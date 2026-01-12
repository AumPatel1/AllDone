"""
Resume exporter module.
Exports resume data to various formats (DOCX, PDF, etc.)
"""

import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_docx(resume_data: Dict[str, Any], output_path: str) -> None:
    """
    Builds a DOCX document from resume data.
    
    Args:
        resume_data: Dictionary containing resume information
        output_path: Path where the DOCX file should be saved
        
    Raises:
        ValueError: If resume_data is invalid
        IOError: If file cannot be saved
    """
    if not isinstance(resume_data, dict):
        raise ValueError("resume_data must be a dictionary")
    
    if not output_path:
        raise ValueError("output_path cannot be empty")
    
    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    try:
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Name (header)
        name = resume_data.get("name", "Resume")
        if name:
            heading = doc.add_heading(name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_info = []
        if resume_data.get("email"):
            contact_info.append(resume_data["email"])
        if resume_data.get("mobile_number"):
            contact_info.append(resume_data["mobile_number"])
        
        if contact_info:
            contact_para = doc.add_paragraph(" | ".join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing

        # Summary
        summary = resume_data.get("summary", "")
        if summary:
            doc.add_heading("Summary", level=2)
            if isinstance(summary, list):
                summary_text = " ".join(summary)
            else:
                summary_text = str(summary)
            doc.add_paragraph(summary_text)
            doc.add_paragraph()  # Spacing

        # Skills
        skills = resume_data.get("skills", [])
        if skills:
            doc.add_heading("Skills", level=2)
            if isinstance(skills, list):
                skills_text = ", ".join(str(skill) for skill in skills)
            elif isinstance(skills, dict):
                # Handle skills as dictionary with categories
                skills_list = []
                for category, skill_list in skills.items():
                    if isinstance(skill_list, list):
                        skills_list.extend(skill_list)
                    else:
                        skills_list.append(str(skill_list))
                skills_text = ", ".join(skills_list)
            else:
                skills_text = str(skills)
            doc.add_paragraph(skills_text)
            doc.add_paragraph()  # Spacing

        # Experience
        experience = resume_data.get("experience", [])
        if experience:
            doc.add_heading("Experience", level=2)
            
            if isinstance(experience, list) and len(experience) > 0:
                if isinstance(experience[0], dict):
                    # Experience is a list of role dictionaries
                    for role in experience:
                        title = role.get("title", "Unknown Position")
                        company = role.get("company", "Unknown Company")
                        doc.add_heading(f'{title} – {company}', level=3)
                        
                        bullets = role.get("bullets", [])
                        if isinstance(bullets, list):
                            for bullet in bullets:
                                doc.add_paragraph(str(bullet), style="List Bullet")
                        else:
                            doc.add_paragraph(str(bullets), style="List Bullet")
                else:
                    # Experience is a list of strings
                    for exp_item in experience:
                        doc.add_paragraph(str(exp_item), style="List Bullet")
            
            doc.add_paragraph()  # Spacing

        # Projects
        projects = resume_data.get("projects", [])
        if projects:
            doc.add_heading("Projects", level=2)
            if isinstance(projects, list):
                for project in projects:
                    doc.add_paragraph(str(project), style="List Bullet")
            else:
                doc.add_paragraph(str(projects))
            doc.add_paragraph()  # Spacing

        # Education
        education = resume_data.get("education", [])
        if education:
            doc.add_heading("Education", level=2)
            if isinstance(education, list):
                for edu_item in education:
                    doc.add_paragraph(str(edu_item), style="List Bullet")
            else:
                doc.add_paragraph(str(education))

        # Save document
        doc.save(output_path)
        print(f"Resume exported successfully to: {output_path}")
        
    except Exception as e:
        raise IOError(f"Failed to save DOCX file to {output_path}: {e}") from e
